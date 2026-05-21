import fitz  # PyMuPDF — для PDF файлов
from docx import Document  # python-docx — для Word файлов
import openpyxl  # для Excel файлов
import os
from translator import translate_text


def handle_pdf(input_path: str, source_lang: str, target_lang: str) -> str:
    """
    Открываем PDF, переводим каждый текстовый блок,
    сохраняем новый PDF рядом с оригиналом.
    """
    doc = fitz.open(input_path)
    output_path = input_path.replace(".pdf", "_translated.pdf")

    for page in doc:
        # get_text("blocks") возвращает список блоков текста на странице
        blocks = page.get_text("blocks")
        for block in blocks:
            if block[6] == 0:  # 0 = текстовый блок (не картинка)
                original = block[4].strip()
                if original:
                    translated = translate_text(original, source_lang, target_lang)
                    rect = fitz.Rect(block[:4])  # координаты блока на странице
                    page.add_redact_annot(rect)   # закрашиваем оригинал
                    page.apply_redactions()
                    page.insert_text(rect.tl, translated, fontsize=10)  # вставляем перевод

    doc.save(output_path)
    doc.close()
    return output_path


def handle_docx(input_path: str, source_lang: str, target_lang: str) -> str:
    """
    Открываем Word документ, переводим каждый абзац и ячейки таблиц,
    сохраняем новый файл.
    """
    doc = Document(input_path)
    output_path = input_path.replace(".docx", "_translated.docx")

    # Переводим обычные абзацы
    for para in doc.paragraphs:
        if para.text.strip():
            translated = translate_text(para.text, source_lang, target_lang)
            # Очищаем все runs (части абзаца) и пишем перевод в первый
            for run in para.runs:
                run.text = ""
            if para.runs:
                para.runs[0].text = translated

    # Переводим текст внутри таблиц
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    translated = translate_text(cell.text, source_lang, target_lang)
                    if cell.paragraphs[0].runs:
                        cell.paragraphs[0].runs[0].text = translated

    doc.save(output_path)
    return output_path


def handle_xlsx(input_path: str, source_lang: str, target_lang: str) -> str:
    """
    Открываем Excel файл, переводим все текстовые ячейки,
    сохраняем новый файл.
    """
    wb = openpyxl.load_workbook(input_path)
    output_path = input_path.replace(".xlsx", "_translated.xlsx")

    for sheet in wb.worksheets:
        for row in sheet.iter_rows():
            for cell in row:
                # Переводим только строки, числа и формулы не трогаем
                if isinstance(cell.value, str) and cell.value.strip():
                    cell.value = translate_text(cell.value, source_lang, target_lang)

    wb.save(output_path)
    return output_path
